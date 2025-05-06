
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import unicodedata

def sanitize_filename(name):
    return name.replace(" ", "_").replace("/", "-")

def normalize(text):
    return ''.join(c for c in unicodedata.normalize('NFD', str(text)) if unicodedata.category(c) != 'Mn').lower().replace('_', ' ').replace('-', ' ')

def detect_column(columns, keyword):
    keyword_norm = normalize(keyword)
    for col in columns:
        if keyword_norm in normalize(col):
            return col
    return None

def load_rdv_data(path, jour_debut, jour_fin, mois, annee):
    df = pd.read_excel(path)
    col_annee = detect_column(df.columns, "annee")
    col_mois = detect_column(df.columns, "mois")
    col_jour = detect_column(df.columns, "jour")
    col_com = detect_column(df.columns, "commercial")

    if not all([col_annee, col_mois, col_jour, col_com]):
        return {}

    df['date'] = pd.to_datetime(dict(
        year=df[col_annee],
        month=df[col_mois],
        day=df[col_jour]
    ))

    df_filtre = df[
        (df[col_annee] == annee) &
        (df[col_mois] == mois) &
        (df[col_jour] >= jour_debut) &
        (df[col_jour] <= jour_fin)
    ]

    if df_filtre.empty:
        return {}

    return dict(tuple(df_filtre.groupby(col_com)))

def creer_rapport_rdv(df, commercial, jour_debut, jour_fin, mois, annee, output_dir, logo_path=None):
    doc = Document()

    # Page de garde
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    doc._body.clear_content()

    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(2))
        p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    titre = doc.add_paragraph()
    titre_run = titre.add_run("RAPPORT RDV COMMERCIAL")
    titre_run.bold = True
    titre_run.font.size = Pt(24)
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    mois_nom = datetime(annee, mois, 1).strftime('%B')
    info = doc.add_paragraph()
    info.add_run("\nCommercial : ").bold = True
    info.add_run(commercial + "\n")
    info.add_run("Période : ").bold = True
    info.add_run(f"du {jour_debut} au {jour_fin} {mois_nom} {annee}")

    doc.add_page_break()

    # En-tête
    header = doc.sections[0].header
    para = header.paragraphs[0]
    if logo_path and os.path.exists(logo_path):
        para.add_run().add_picture(logo_path, width=Inches(1.5))
    para.add_run(f"   Compte rendu du {jour_debut} au {jour_fin} {mois_nom} {annee} – RDV de {commercial}")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Table des rendez-vous
    doc.add_heading("Rendez-vous", level=1)
    if df.empty:
        doc.add_paragraph("Aucun rendez-vous trouvé pour cette période.", style="Intense Quote")
    else:
        col_date = detect_column(df.columns, "date")
        col_raison = detect_column(df.columns, "raison")
        col_adresse = detect_column(df.columns, "adresse")

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'
        headers = ['Date', 'Raison du RDV', 'Adresse']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade = OxmlElement('w:shd')
            shade.set(qn('w:fill'), 'D9E1F2')
            cell._tc.get_or_add_tcPr().append(shade)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            cells[0].text = row[col_date].strftime("%d/%m/%Y") if pd.notnull(row[col_date]) else ""
            cells[1].text = str(row[col_raison]) if col_raison else ""
            cells[2].text = str(row[col_adresse]) if col_adresse else ""

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_dir}/RDV_{sanitize_filename(commercial)}_{mois:02d}_{annee}.docx"
    doc.save(filename)
    return filename
